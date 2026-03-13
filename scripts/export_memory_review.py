from __future__ import annotations

import argparse
import json
import sqlite3
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DB_PATH = ROOT / "data" / "doctor_memory.db"
DEFAULT_OUTPUT_DIR = ROOT / "exports"
EXCLUDED_KINDS = {"system_marker"}

KIND_LABELS = {
    "academic_role": "学术身份",
    "clinical_strength": "临床强项",
    "diagnostic_principle": "诊断原则",
    "doctor_thought": "医生思路",
    "escalation_rule": "升级转诊规则",
    "follow_up_rule": "随访规则",
    "hard_boundary": "硬性边界",
    "knowledge_seed": "基础知识种子",
    "patient_education": "患者宣教",
    "public_bio": "公开职业信息",
    "research_direction": "研究方向",
    "risk_rule": "风险规则",
    "specialty_focus": "专科重点方向",
    "thinking_style": "表达与思维风格",
    "treatment_principle": "治疗原则",
}

KIND_NOTES = {
    "academic_role": "用于确认医生的公开身份与学术角色。",
    "clinical_strength": "用于界定公开可见的专科强项与擅长方向。",
    "diagnostic_principle": "用于说明面对症状或疾病时，优先如何判断与检查。",
    "doctor_thought": "更接近医生本人的经验判断、取舍逻辑和临床关注点。",
    "escalation_rule": "用于标记何时应升级处理、转诊或尽快线下就医。",
    "follow_up_rule": "用于记录复诊、复查、疗程观察时的常见跟进思路。",
    "hard_boundary": "明确哪些事情 AI 助手不能做、不能说。",
    "knowledge_seed": "来自基础资料的种子知识，可作为后续进一步提炼的原始素材。",
    "patient_education": "用于向患者解释疾病、检查、治疗与风险。",
    "public_bio": "公开渠道可验证的职业背景信息。",
    "research_direction": "公开可见的研究与长期关注方向。",
    "risk_rule": "用于提醒危险信号、并发症或高风险情况。",
    "specialty_focus": "专科内更聚焦的方向。",
    "thinking_style": "更接近你爸的说话口吻、沟通偏好和思考方式。",
    "treatment_principle": "用于说明治疗取向、保守与积极处理边界、方案选择原则。",
}

DOC_FONT = "Microsoft YaHei"
PDF_FONT_REGULAR = "DoctorMemoryCN"
PDF_FONT_BOLD = "DoctorMemoryCNBold"
PDF_FONT_ITALIC = "DoctorMemoryCN"


def ensure_pdf_fonts() -> None:
    regular_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")

    if PDF_FONT_REGULAR not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_FONT_REGULAR, str(regular_path), subfontIndex=0))
    if PDF_FONT_BOLD not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(bold_path), subfontIndex=0))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export doctor memory DB to a readable review document.")
    parser.add_argument("--db", default=str(DEFAULT_DB_PATH), help="Path to SQLite database.")
    parser.add_argument("--out-dir", default=str(DEFAULT_OUTPUT_DIR), help="Directory for export files.")
    parser.add_argument(
        "--basename",
        default=f"doctor-memory-review-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        help="Base filename without extension.",
    )
    parser.add_argument(
        "--mode",
        choices=["full", "slim"],
        default="full",
        help="slim mode keeps only title, content and source in review documents.",
    )
    parser.add_argument(
        "--formats",
        default="md,json,docx,pdf",
        help="Comma-separated output formats from: md,json,docx,pdf",
    )
    return parser.parse_args()


def load_rows(db_path: Path) -> list[dict]:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    try:
        rows = conn.execute(
            """
            SELECT id, kind, title, content, tags_json, source, importance, created_at, updated_at
            FROM doctor_memory_entries
            ORDER BY kind ASC, importance DESC, updated_at DESC, id ASC
            """
        ).fetchall()
    finally:
        conn.close()

    result = []
    for row in rows:
        try:
            tags = json.loads(row["tags_json"] or "[]")
        except json.JSONDecodeError:
            tags = []
        result.append(
            {
                "id": row["id"],
                "kind": row["kind"],
                "title": row["title"],
                "content": row["content"],
                "tags": tags,
                "source": row["source"],
                "importance": row["importance"],
                "created_at": row["created_at"],
                "updated_at": row["updated_at"],
            }
        )
    return [row for row in result if row["kind"] not in EXCLUDED_KINDS]


def render_markdown(rows: list[dict], db_path: Path) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    by_kind: dict[str, list[dict]] = {}
    for row in rows:
        by_kind.setdefault(row["kind"], []).append(row)

    counter = Counter(row["kind"] for row in rows)

    lines: list[str] = []
    lines.append("# 李勇医生资料库导出审阅版")
    lines.append("")
    lines.append("这份文档按资料类别整理，目的是方便人工审阅当前数据库内容是否合适、是否有表达不准确、边界不稳妥、风格不像本人等问题。")
    lines.append("")
    lines.append("## 导出信息")
    lines.append("")
    lines.append(f"- 导出时间：`{now}`")
    lines.append(f"- 数据库路径：`{db_path}`")
    lines.append(f"- 总条数：`{len(rows)}`")
    lines.append("")
    lines.append("## 分类概览")
    lines.append("")
    for kind, count in sorted(counter.items(), key=lambda item: (-item[1], KIND_LABELS.get(item[0], item[0]))):
        label = KIND_LABELS.get(kind, kind)
        lines.append(f"- `{label}`：`{count}` 条")
    lines.append("")
    lines.append("## 审阅建议")
    lines.append("")
    lines.append("- 重点看 `医生思路 / 表达与思维风格 / 患者宣教 / 风险规则` 是否真的像李勇医生本人。")
    lines.append("- 重点删改过于模板化、太像 AI 写的句子，或与本人真实说法不一致的内容。")
    lines.append("- 重点确认医疗边界是否保守、稳妥，避免出现替代面诊、替代检查、直接诊断处方的表述。")
    lines.append("")

    for kind in sorted(by_kind.keys(), key=lambda item: (KIND_LABELS.get(item, item))):
        label = KIND_LABELS.get(kind, kind)
        note = KIND_NOTES.get(kind, "")
        entries = by_kind[kind]

        lines.append(f"## {label}")
        lines.append("")
        if note:
            lines.append(note)
            lines.append("")
        lines.append(f"- 条数：`{len(entries)}`")
        lines.append("")

        for idx, entry in enumerate(entries, start=1):
            lines.append(f"### {idx}. {entry['title']}")
            lines.append("")
            lines.append(f"- ID：`{entry['id']}`")
            lines.append(f"- 重要度：`{entry['importance']}`")
            lines.append(f"- 来源：`{entry['source']}`")
            if entry["tags"]:
                tag_text = " / ".join(str(tag) for tag in entry["tags"])
                lines.append(f"- 标签：`{tag_text}`")
            if entry["updated_at"]:
                lines.append(f"- 更新时间：`{entry['updated_at']}`")
            lines.append("- 内容：")
            lines.append("")
            for paragraph in (entry["content"] or "").splitlines():
                paragraph = paragraph.rstrip()
                if paragraph:
                    lines.append(f"  {paragraph}")
                else:
                    lines.append("")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def render_json(rows: list[dict]) -> str:
    payload = {
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "row_count": len(rows),
        "entries": rows,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def export_docx(rows: list[dict], db_path: Path, out_path: Path, *, mode: str = "full") -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = DOC_FONT
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("李勇医生资料库精简审阅版" if mode == "slim" else "李勇医生资料库导出审阅版")
    run.bold = True
    run.font.name = DOC_FONT
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "按资料类别整理，仅保留标题、内容和来源，便于医生快速审阅"
        if mode == "slim"
        else "按资料类别整理，便于人工审阅与修改"
    )
    subtitle_run.font.name = DOC_FONT
    subtitle_run.font.size = Pt(10.5)

    document.add_paragraph("")
    document.add_heading("导出信息", level=1)
    info_items = [
        f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"数据库路径：{db_path}",
        f"总条数：{len(rows)}",
    ]
    for item in info_items:
        document.add_paragraph(item, style="List Bullet")

    counter = Counter(row["kind"] for row in rows)
    document.add_heading("分类概览", level=1)
    for kind, count in sorted(counter.items(), key=lambda item: (-item[1], KIND_LABELS.get(item[0], item[0]))):
        document.add_paragraph(f"{KIND_LABELS.get(kind, kind)}：{count} 条", style="List Bullet")

    document.add_heading("审阅建议", level=1)
    review_tips = [
        "重点看“医生思路 / 表达与思维风格 / 患者宣教 / 风险规则”是否真的像李勇医生本人。",
        "重点删改过于模板化、太像 AI 写的句子，或与本人真实说法不一致的内容。",
        "重点确认医疗边界是否保守、稳妥，避免出现替代面诊、替代检查、直接诊断处方的表述。",
    ]
    for tip in review_tips:
        document.add_paragraph(tip, style="List Bullet")

    by_kind: dict[str, list[dict]] = {}
    for row in rows:
        by_kind.setdefault(row["kind"], []).append(row)

    for kind in sorted(by_kind.keys(), key=lambda item: KIND_LABELS.get(item, item)):
        label = KIND_LABELS.get(kind, kind)
        note = KIND_NOTES.get(kind, "")
        entries = by_kind[kind]

        document.add_page_break()
        document.add_heading(label, level=1)
        if note:
            document.add_paragraph(note)
        document.add_paragraph(f"条数：{len(entries)}")

        for idx, entry in enumerate(entries, start=1):
            document.add_heading(f"{idx}. {entry['title']}", level=2)
            if mode == "slim":
                meta_lines = [f"来源：{entry['source']}"]
            else:
                meta_lines = [
                    f"ID：{entry['id']}",
                    f"重要度：{entry['importance']}",
                    f"来源：{entry['source']}",
                ]
                if entry["tags"]:
                    meta_lines.append(f"标签：{' / '.join(str(tag) for tag in entry['tags'])}")
                if entry["updated_at"]:
                    meta_lines.append(f"更新时间：{entry['updated_at']}")
            for meta in meta_lines:
                p = document.add_paragraph()
                p.add_run(meta).italic = True
            document.add_paragraph(entry["content"] or "")

    document.save(out_path)


def export_pdf(rows: list[dict], db_path: Path, out_path: Path, *, mode: str = "full") -> None:
    ensure_pdf_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=PDF_FONT_BOLD,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1f2937"),
        spaceAfter=8 if mode == "slim" else 10,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleCN",
        parent=styles["Normal"],
        fontName=PDF_FONT_REGULAR,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4b5563"),
        spaceAfter=12 if mode == "slim" else 18,
    )
    h1 = ParagraphStyle(
        "H1CN",
        parent=styles["Heading1"],
        fontName=PDF_FONT_BOLD,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#111827"),
        spaceBefore=8 if mode == "slim" else 12,
        spaceAfter=6 if mode == "slim" else 8,
    )
    h2 = ParagraphStyle(
        "H2CN",
        parent=styles["Heading2"],
        fontName=PDF_FONT_BOLD,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#1f2937"),
        spaceBefore=7 if mode == "slim" else 10,
        spaceAfter=4 if mode == "slim" else 6,
    )
    body = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=PDF_FONT_REGULAR,
        fontSize=9.0 if mode == "slim" else 9.5,
        leading=13 if mode == "slim" else 15,
        textColor=colors.black,
        spaceAfter=2 if mode == "slim" else 4,
    )
    meta = ParagraphStyle(
        "MetaCN",
        parent=body,
        fontName=PDF_FONT_ITALIC,
        fontSize=8.2,
        leading=11,
        textColor=colors.HexColor("#4b5563"),
        leftIndent=6,
    )

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    story = [
        Paragraph("李勇医生资料库精简审阅版" if mode == "slim" else "李勇医生资料库导出审阅版", title_style),
        Paragraph(
            "按资料类别整理，仅保留标题、内容和来源，便于医生快速审阅"
            if mode == "slim"
            else "按资料类别整理，便于人工审阅与修改",
            subtitle_style,
        ),
        Paragraph("导出信息", h1),
        Paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body),
        Paragraph(f"数据库路径：{db_path}", body),
        Paragraph(f"总条数：{len(rows)}", body),
        Spacer(1, 8),
    ]

    counter = Counter(row["kind"] for row in rows)
    story.append(Paragraph("分类概览", h1))
    for kind, count in sorted(counter.items(), key=lambda item: (-item[1], KIND_LABELS.get(item[0], item[0]))):
        story.append(Paragraph(f"- {KIND_LABELS.get(kind, kind)}：{count} 条", body))

    story.append(Paragraph("审阅建议", h1))
    for tip in [
        "重点看“医生思路 / 表达与思维风格 / 患者宣教 / 风险规则”是否真的像李勇医生本人。",
        "重点删改过于模板化、太像 AI 写的句子，或与本人真实说法不一致的内容。",
        "重点确认医疗边界是否保守、稳妥，避免出现替代面诊、替代检查、直接诊断处方的表述。",
    ]:
        story.append(Paragraph(f"- {tip}", body))

    by_kind: dict[str, list[dict]] = {}
    for row in rows:
        by_kind.setdefault(row["kind"], []).append(row)

    for kind in sorted(by_kind.keys(), key=lambda item: KIND_LABELS.get(item, item)):
        label = KIND_LABELS.get(kind, kind)
        note = KIND_NOTES.get(kind, "")
        entries = by_kind[kind]

        if mode != "slim":
            story.append(PageBreak())
        story.append(Paragraph(label, h1))
        if note:
            story.append(Paragraph(note, body))
        story.append(Paragraph(f"条数：{len(entries)}", body))

        for idx, entry in enumerate(entries, start=1):
            story.append(Paragraph(f"{idx}. {entry['title']}", h2))
            story.append(Paragraph(f"来源：{entry['source']}", meta))
            if mode != "slim":
                story.append(Paragraph(f"ID：{entry['id']}", meta))
                story.append(Paragraph(f"重要度：{entry['importance']}", meta))
            if mode != "slim" and entry["tags"]:
                story.append(Paragraph(f"标签：{' / '.join(str(tag) for tag in entry['tags'])}", meta))
            if mode != "slim" and entry["updated_at"]:
                story.append(Paragraph(f"更新时间：{entry['updated_at']}", meta))
            for paragraph in (entry["content"] or "").splitlines():
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), body))
            story.append(Spacer(1, 4 if mode == "slim" else 6))

    doc.build(story)


def main() -> None:
    args = parse_args()
    db_path = Path(args.db)
    if not db_path.exists():
        raise SystemExit(f"Database not found: {db_path}")

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    rows = load_rows(db_path)
    md_path = out_dir / f"{args.basename}.md"
    json_path = out_dir / f"{args.basename}.json"
    docx_path = out_dir / f"{args.basename}.docx"
    pdf_path = out_dir / f"{args.basename}.pdf"
    formats = {item.strip().lower() for item in args.formats.split(",") if item.strip()}

    if "md" in formats:
        md_path.write_text(render_markdown(rows, db_path), encoding="utf-8")
    if "json" in formats:
        json_path.write_text(render_json(rows), encoding="utf-8")
    if "docx" in formats:
        export_docx(rows, db_path, docx_path, mode=args.mode)
    if "pdf" in formats:
        export_pdf(rows, db_path, pdf_path, mode=args.mode)

    print(f"Rows: {len(rows)}")
    if "md" in formats:
        print(f"Markdown: {md_path}")
    if "json" in formats:
        print(f"JSON: {json_path}")
    if "docx" in formats:
        print(f"DOCX: {docx_path}")
    if "pdf" in formats:
        print(f"PDF: {pdf_path}")


if __name__ == "__main__":
    main()
